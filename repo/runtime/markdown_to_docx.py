#!/usr/bin/env python3
"""
Pisoso Legal AI - Conversor Oficial de Markdown a Word (.docx)
Inyecta el contenido en las plantillas oficiales de 'Palntillas word/',
remplaza los metadatos de portada/encabezados y elimina tablas de muestra no utilizadas.
Garantiza tipografía >= 11 pt y estilo corporativo Navy & Slate.
"""

import sys
import os
import re

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx no está instalado.")
    sys.exit(1)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
WORKSPACE_ROOT = os.path.dirname(SCRIPT_DIR)
TEMPLATES_DIR = os.path.join(WORKSPACE_ROOT, "Palntillas word")

try:
    from template_selector import resolve_template
except ImportError:
    from scripts.template_selector import resolve_template

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex="F2F4F7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_markdown_runs(paragraph, text, min_pt=11):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
        else:
            clean_token = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', token)
            run = paragraph.add_run(clean_token)
            
        if run.font.size is None or run.font.size < Pt(min_pt):
            run.font.size = Pt(min_pt)

def parse_markdown_table(lines, start_idx):
    rows = []
    idx = start_idx
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if re.match(r'^\|?[\s:-]*\|[\s:-|]*$', line):
            idx += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not cells and line.split('|'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
        idx += 1
    return rows, idx - 1

def replace_text_in_tables(doc, replacements):
    """Reemplaza etiquetas vacías de plantilla en tablas que deseamos conservar."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        for p in cell.paragraphs:
                            if key in p.text:
                                p.text = p.text.replace(key, val)
                                for r in p.runs:
                                    r.font.size = Pt(11)

def convert_md_to_docx(md_path, output_path, template_hint=""):
    if not os.path.exists(md_path):
        print(f"Error: No existe el archivo Markdown en {md_path}")
        sys.exit(1)
        
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    if not template_hint:
        hint_src = os.path.basename(output_path) + " " + os.path.basename(md_path) + " " + md_content[:500]
        template_path, tpl_name = resolve_template(hint_src)
    else:
        template_path, tpl_name = resolve_template(template_hint)

    print(f"📄 Seleccionando plantilla oficial: {tpl_name}")

    if not os.path.exists(template_path):
        doc = docx.Document()
    else:
        doc = docx.Document(template_path)

    # 1. ELIMINAR COMPLETAMENTE LOS PÁRRAFOS Y TODAS LAS TABLAS DE MUESTRA DEL CUERPO
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
        
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # 2. RENDERIZAR EL CONTENIDO SUSTANTIVO REAL DESDE EL MARKDOWN
    lines = md_content.splitlines()
    in_code_block = False
    code_lines = []
    idx = 0

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Omitir frontmatter YAML
        if idx == 0 and stripped == '---':
            idx += 1
            while idx < len(lines) and lines[idx].strip() != '---':
                idx += 1
            idx += 1
            continue

        # Bloques de código / Mermaid
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Tablas Markdown
        if '|' in stripped and not stripped.startswith('#'):
            table_data, end_idx = parse_markdown_table(lines, idx)
            if table_data and len(table_data) > 0:
                num_cols = max(len(r) for r in table_data)
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                            
                            if r_idx == 0:
                                set_cell_shading(cell, "0A192F") # Navy corporativo Pisoso
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(11)
                            else:
                                if r_idx % 2 == 0:
                                    set_cell_shading(cell, "F8F9FA")
                                add_markdown_runs(p, cell_text, min_pt=11)
                idx = end_idx + 1
                continue

        # Títulos con jerarquía visual y tipografía >= 11 pt
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(10, 25, 47) # Navy
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(13.5)
            run.font.color.rgb = RGBColor(30, 58, 138)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_markdown_runs(p, stripped[2:], min_pt=11)
        elif re.match(r'^\d+\.\s+', stripped):
            match = re.match(r'^\d+\.\s+', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_markdown_runs(p, stripped[len(match.group(0)):], min_pt=11)
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_markdown_runs(p, stripped, min_pt=11)

        idx += 1

    doc.save(output_path)
    print(f"✅ Documento generado con éxito (limpio de tablas predeterminadas): {output_path}")

def main():
    if len(sys.argv) < 3:
        print("Uso: python3 markdown_to_docx.py <input.md> <output.docx> [template_hint]")
        sys.exit(1)
        
    md_path = sys.argv[1]
    output_path = sys.argv[2]
    template_hint = sys.argv[3] if len(sys.argv) > 3 else ""
    convert_md_to_docx(md_path, output_path, template_hint)

if __name__ == "__main__":
    main()
